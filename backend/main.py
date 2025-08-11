# main.py
# Agentic Framework Backend for the AI Business Analyst - FINAL

from flask import Flask, request, jsonify, redirect, send_file, send_from_directory
from flask_cors import CORS
import os
from io import BytesIO
import requests
import json
import uuid
import base64
import io
import re
from datetime import datetime
from werkzeug.utils import secure_filename

# --- File Parsing Libraries ---
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
import PyPDF2
import docx

# --- Azure SDKs ---
from azure.communication.email import EmailClient

# --- Database imports ---
from database import (
    init_db, get_db, save_document_to_db, save_document_to_db_direct, save_analysis_to_db,
    add_to_vector_db, search_vector_db, Document, Analysis,
    save_approval_to_db, get_approval_from_db, update_approval_in_db,
    get_all_documents_from_db, get_all_documents_from_db_direct, get_all_analyses_from_db, 
    check_document_exists_by_name, check_document_exists_by_name_direct, get_document_by_id, get_analysis_by_id_from_db
)

# --- Configuration imports ---
from config import (
    GEMINI_API_URL, ACS_CONNECTION_STRING, ACS_SENDER_ADDRESS,
    APPROVAL_RECIPIENT_EMAIL, BACKEND_BASE_URL, ADO_ORGANIZATION_URL,
    ADO_PROJECT_NAME, ADO_PAT, DATABASE_URL, GEMINI_API_KEY
)

# Prevent Flask from auto-loading .env with hardcoded UTF-8
os.environ.setdefault("FLASK_SKIP_DOTENV", "1")

# Initialize the Flask application
app = Flask(__name__)

# Initialize database only if configured to do so
try:
    init_db()
except Exception as _e:
    # Do not block startup; table creation can be done explicitly later
    print(f"ℹ️ Skipping DB init at startup: {_e}")

# --- Configuration ---
CORS(app, resources={r"/api/*": {"origins": "*"}})

# --- In-memory storage for approvals (fallback only) ---
approval_statuses = {}

# --- Token Consumption Tracking ---
token_consumption_logs = []

def log_token_consumption(stage, tokens_used, model_used="gemini-pro", details=None):
    """Log token consumption for each stage with detailed information"""
    log_entry = {
        "timestamp": datetime.now().isoformat(),
        "stage": stage,
        "tokens_used": tokens_used,
        "model_used": model_used,
        "details": details or {}
    }
    token_consumption_logs.append(log_entry)
    
    # Save to file for persistence (use tmp on serverless)
    try:
        log_dir = os.getenv('TMP_DIR', '/tmp')
        os.makedirs(log_dir, exist_ok=True)
        with open(os.path.join(log_dir, "token_logs.json"), "w") as f:
            json.dump(token_consumption_logs, f, indent=2)
    except Exception as e:
        print(f"Warning: Could not save token logs: {e}")
    
    # Print detailed log to console
    print(f"🔍 TOKEN LOG: {stage.upper()}")
    print(f"   📊 Tokens Used: {tokens_used:,}")
    print(f"   🤖 Model: {model_used}")
    print(f"   ⏰ Timestamp: {log_entry['timestamp']}")
    if details:
        print(f"   📝 Details: {details}")
    print(f"   📈 Total Tokens So Far: {sum(log['tokens_used'] for log in token_consumption_logs):,}")
    print("─" * 50)

@app.route("/api/token_logs", methods=['GET'])
def get_token_logs():
    """Get token consumption logs"""
    return jsonify(token_consumption_logs)

# --- Document Management ---
@app.route("/api/upload_document", methods=['POST'])
def upload_document():
    """Upload and store a document"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        print(f"📥 [Upload] Started for file: {file.filename}")

        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        if not file.filename.lower().endswith(('.pdf', '.docx')):
            return jsonify({"error": "Only PDF and DOCX files are allowed"}), 400

        print(f"DEBUG: Starting document upload for {file.filename}")

        # Check if document with same name already exists
        print(f"🔍 [Upload] Checking if file '{file.filename}' already exists...")

        if check_document_exists_by_name_direct(file.filename):
            print(f"⚠️ [Upload] File already exists: {file.filename}")

            print(f"DEBUG: Document with name '{file.filename}' already exists, skipping upload")
            return jsonify({
                "error": f"Document with name '{file.filename}' already exists. Please use a different name.",
                "duplicate": True
            }), 409  # Conflict status code

        # Read file content once and store it
        file_content_bytes = file.read()
        file_size = len(file_content_bytes)
        print(f"DEBUG: File size: {file_size} bytes")
        
        # Create uploads directory (use serverless tmp)
        uploads_dir = os.getenv('UPLOADS_DIR', '/tmp/uploads')
        if not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir)
            print(f"📁 Created uploads directory: {uploads_dir}")
        
        # Extract content from file
        file_content, images, error = agent_extract_content(io.BytesIO(file_content_bytes), file.filename)
        if error:
            print(f"DEBUG: Content extraction failed: {error}")
            return jsonify({"error": f"Failed to extract content: {error}"}), 500

        print(f"DEBUG: Content extracted successfully, length: {len(file_content)} characters")

        # Create document record
        doc_id = str(uuid.uuid4())
        # Save file to disk after id is created
        safe_name = secure_filename(file.filename)
        file_path = os.path.join(uploads_dir, f"{doc_id}_{safe_name}")
        with open(file_path, 'wb') as f:
            f.write(file_content_bytes)
        print(f"💾 [File] Saved file to disk: {file_path}")
        document_data = {
            "id": doc_id,
            "name": file.filename,
            "upload_date": datetime.now().isoformat(),
            "file_type": file.filename.split('.')[-1].lower(),
            "size": file_size,
            "status": "uploaded"
        }
        
        print(f"DEBUG: Document data created with ID: {doc_id}")
        
        # Save to database using direct psycopg2 approach (working method)
        try:
            print(f"💾 [DB] Saving document to database using direct psycopg2: {file.filename}")
            save_document_to_db_direct(document_data, file_path, file_content)
            print(f"✅ [DB] Document stored with ID: {doc_id}")

            # Add to vector database
            print(f"📡 [VectorDB] Adding to vector DB...")
            add_to_vector_db(
                content=file_content,
                meta={
                    "id": doc_id,
                    "name": file.filename,
                    "type": "document",
                    "upload_date": document_data["upload_date"]
                },
                collection_name="documents"
            )
            print(f"✅ [VectorDB] Added to vector DB")

        except Exception as db_error:
            print(f"❌ [DB] Failed to save document: {db_error}")
            import traceback
            traceback.print_exc()
            raise db_error

        # no db session to close here; using direct psycopg2 path above


        
        print(f"DEBUG: Document upload completed successfully for {file.filename}")
        return jsonify({
            "success": True,
            "message": "Document uploaded successfully",
            "document_id": doc_id,
            "filename": file.filename,
            "file_path": file_path,
            "size": file_size,
            "content_length": len(file_content) if file_content else 0
        }), 201
        
    except Exception as e:
        print(f"❌ Error uploading document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to upload document: {str(e)}"}), 500

@app.route("/api/test_db", methods=['GET'])
def test_database():
    """Test database connectivity and document count"""
    try:
        db = next(get_db())
        try:
            # Test document count
            doc_count = db.query(Document).count()
            analysis_count = db.query(Analysis).count()
            
            return jsonify({
                "success": True,
                "message": "Database connection successful",
                "document_count": doc_count,
                "analysis_count": analysis_count,
                "database_url": DATABASE_URL.split('@')[0] + "@[HIDDEN]"  # Hide credentials
            })
        finally:
            db.close()
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"Database connection failed: {str(e)}",
            "error": str(e)
        }), 500

@app.route("/api/test_create_doc", methods=['POST'])
def test_create_document():
    """Test endpoint to manually create a document in the database"""
    try:
        data = request.get_json()
        test_content = data.get('content', 'Test document content')
        test_filename = data.get('filename', 'test_document.txt')
        
        doc_id = str(uuid.uuid4())
        document_data = {
            "id": doc_id,
            "name": test_filename,
            "uploadDate": datetime.now().isoformat(),
            "fileType": "txt",
            "size": len(test_content),
            "status": "uploaded"
        }
        
        print(f"DEBUG: Testing document creation with ID: {doc_id}")
        
        db = next(get_db())
        try:
            save_document_to_db(db, document_data, f"test_uploads/{doc_id}_{test_filename}", test_content)
            
            return jsonify({
                "success": True,
                "message": "Test document created successfully",
                "document_id": doc_id,
                "document_data": document_data
            })
        finally:
            db.close()
            
    except Exception as e:
        print(f"Error in test document creation: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to create test document: {str(e)}"}), 500

@app.route("/api/documents", methods=['GET'])
def get_documents():
    """Get all documents from database using direct psycopg2"""
    try:
        documents = get_all_documents_from_db_direct()
        print(f"DEBUG: Retrieved {len(documents)} documents from database using direct psycopg2")
        return jsonify(documents)
    except Exception as e:
        print(f"❌ Error retrieving documents: {e}")
        return jsonify({"error": f"Failed to retrieve documents: {str(e)}"}), 500

@app.route("/api/documents/files", methods=['GET'])
def list_uploaded_files():
    """List all files in the uploads directory"""
    try:
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            return jsonify({"files": [], "message": "Uploads directory does not exist"})
        
        files = []
        for filename in os.listdir(uploads_dir):
            file_path = os.path.join(uploads_dir, filename)
            if os.path.isfile(file_path):
                file_size = os.path.getsize(file_path)
                files.append({
                    "name": filename,
                    "size": file_size,
                    "path": file_path,
                    "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
                })
        
        return jsonify({
            "files": files,
            "count": len(files),
            "directory": uploads_dir
        })
    except Exception as e:
        return jsonify({"error": f"Failed to list files: {str(e)}"}), 500

@app.route("/api/documents/<doc_id>", methods=['GET'])
def get_document(doc_id):
    """Get a specific document by ID"""
    try:
        db = next(get_db())
        try:
            doc = get_document_by_id(db, doc_id)
            if doc:
                return jsonify(doc)
            else:
                return jsonify({"error": "Document not found"}), 404
        finally:
            db.close()
    except Exception as e:
        return jsonify({"error": f"Failed to retrieve document: {str(e)}"}), 500

@app.route("/api/documents/<doc_id>/file", methods=['GET'])
def download_document(doc_id):
    """Download the actual file for a document"""
    try:
        db = next(get_db())
        try:
            doc = get_document_by_id(db, doc_id)
            if doc and os.path.exists(doc['file_path']):
                return send_file(doc['file_path'], as_attachment=True, download_name=doc['name'])
            else:
                return jsonify({"error": "Document file not found"}), 404
        finally:
            db.close()
    except Exception as e:
        return jsonify({"error": f"Failed to download document: {str(e)}"}), 500

# --- Past Analyses Management ---
@app.route("/api/analyses", methods=['GET'])
def get_past_analyses():
    """Get all past analyses from database"""
    try:
        db = next(get_db())
        try:
            analyses = get_all_analyses_from_db(db)
            print(f"DEBUG: Retrieved {len(analyses)} analyses from database")
            return jsonify(analyses)
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error retrieving analyses: {e}")
        return jsonify({"error": f"Failed to retrieve analyses: {str(e)}"}), 500

@app.route("/api/analyses/<analysis_id>", methods=['GET'])
def get_analysis(analysis_id):
    """Get a specific analysis by ID from database"""
    try:
        db = next(get_db())
        try:
            analysis = get_analysis_by_id_from_db(db, analysis_id)
            if analysis:
                print(f"DEBUG: Retrieved analysis '{analysis['title']}' from database")
                return jsonify(analysis)
            else:
                print(f"DEBUG: Analysis with ID '{analysis_id}' not found")
                return jsonify({"error": "Analysis not found"}), 404
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error retrieving analysis: {e}")
        return jsonify({"error": f"Failed to retrieve analysis: {str(e)}"}), 500

def get_analysis_by_id(analysis_id):
    """Get analysis by ID from database"""
    try:
        db = next(get_db())
        try:
            analysis = get_analysis_by_id_from_db(db, analysis_id)
            return analysis
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error getting analysis by ID: {e}")
        return None

def save_analysis_results(results, original_text, filename, user_email=None):
    """Save analysis results to database"""
    try:
        analysis_id = str(uuid.uuid4())
        analysis_data = {
            "id": analysis_id,
            "title": f"Analysis of {filename}",
            "originalText": original_text,
            "results": results,
            "user_email": user_email
        }
        
        db = next(get_db())
        try:
            save_analysis_to_db(db, analysis_data)
            print(f"✅ Analysis saved to database: {analysis_id}")
            return analysis_id
        finally:
            db.close()
            
    except Exception as e:
        print(f"❌ Error saving analysis results: {e}")
        return None

# --- Helper function to convert markdown to DOCX ---
def markdown_to_docx(markdown_content):
    """Convert markdown content to DOCX format with enhanced table support"""
    doc = DocxDocument()
    
    # Split content into lines
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            doc.add_paragraph()
            i += 1
            continue
            
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level == 1:
                doc.add_heading(text, 0)
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            elif level == 4:
                doc.add_heading(text, 3)
            else:
                doc.add_heading(text, 4)
            i += 1
            
        # Handle tables
        elif line.startswith('|') and line.endswith('|'):
            # Start of a table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:  # Need at least header and separator
                # Parse table
                table_data = []
                for table_line in table_lines:
                    if table_line.startswith('|') and table_line.endswith('|'):
                        # Remove outer pipes and split
                        cells = table_line[1:-1].split('|')
                        # Clean up each cell
                        cells = [cell.strip() for cell in cells]
                        table_data.append(cells)
                
                if table_data:
                    # Create table
                    num_rows = len(table_data)
                    num_cols = max(len(row) for row in table_data) if table_data else 0
                    
                    if num_rows > 0 and num_cols > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Table Grid'
                        
                        # Fill table data
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_text in enumerate(row_data):
                                if col_idx < num_cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = cell_text
                                    
                                    # Make header row bold
                                    if row_idx == 0:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
        
        # Handle bold text
        elif '**' in line:
            # Simple bold handling - replace **text** with bold
            p = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Bold text
                    p.add_run(part).bold = True
                else:  # Regular text
                    p.add_run(part)
            i += 1
            
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
            i += 1
            
        # Handle code blocks
        elif line.startswith('```'):
            # Skip code block markers
            i += 1
            # Add code block content as preformatted text
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_line = lines[i]
                p = doc.add_paragraph(code_line)
                # Apply monospace font to code
                for run in p.runs:
                    run.font.name = 'Courier New'
                i += 1
            if i < len(lines) and lines[i].strip().startswith('```'):
                i += 1  # Skip closing ```
                
        # Handle regular paragraphs
        else:
            doc.add_paragraph(line)
            i += 1
    
    return doc

def convert_image_format(image_data, from_mime_type, to_mime_type='image/png'):
    """Convert image from one format to another (enhanced implementation)"""
    try:
        # Try to import PIL for image conversion
        try:
            from PIL import Image, ImageDraw
            import io
            
            # For WMF files, we need special handling
            if from_mime_type.lower() in ['image/x-wmf', 'image/wmf']:
                print(f"DEBUG: WMF format detected - attempting enhanced handling...")
                
                # Method 1: Try direct PIL opening (rarely works for WMF)
                try:
                    image = Image.open(io.BytesIO(image_data))
                    print(f"DEBUG: Successfully opened WMF with PIL")
                except Exception as wmf_error:
                    print(f"DEBUG: PIL cannot handle WMF directly: {wmf_error}")
                    
                    # Method 2: Try to install and use wand (ImageMagick wrapper)
                    try:
                        from wand.image import Image as WandImage
                        print(f"DEBUG: Attempting WMF conversion with ImageMagick...")
                        
                        with WandImage(blob=image_data) as wand_img:
                            # Convert to PNG format
                            wand_img.format = 'png'
                            converted_data = wand_img.make_blob()
                            print(f"DEBUG: Successfully converted WMF to PNG using ImageMagick")
                            return converted_data, 'image/png'
                            
                    except ImportError:
                        print(f"DEBUG: ImageMagick (wand) not available for WMF conversion")
                    except Exception as magick_error:
                        print(f"DEBUG: ImageMagick conversion failed: {magick_error}")
                    
                    # Method 3: Try using Windows-specific libraries (if available)
                    try:
                        import win32com.client
                        print(f"DEBUG: Attempting WMF conversion using Windows COM...")
                        # This would require Windows-specific code
                        print(f"DEBUG: Windows COM conversion not implemented yet")
                    except ImportError:
                        print(f"DEBUG: Windows COM libraries not available")
                    except Exception as com_error:
                        print(f"DEBUG: Windows COM conversion failed: {com_error}")
                    
                    # Method 4: Create a placeholder image
                    print(f"DEBUG: Creating placeholder image for WMF...")
                    try:
                        # Create a simple placeholder image
                        placeholder = Image.new('RGB', (300, 200), color='lightgray')
                        draw = ImageDraw.Draw(placeholder)
                        draw.text((10, 10), "WMF Image", fill='black')
                        draw.text((10, 30), "Format not supported", fill='red')
                        
                        output_buffer = io.BytesIO()
                        placeholder.save(output_buffer, format='PNG')
                        converted_data = output_buffer.getvalue()
                        print(f"DEBUG: Created placeholder image for WMF")
                        return converted_data, 'image/png'
                        
                    except Exception as placeholder_error:
                        print(f"DEBUG: Failed to create placeholder: {placeholder_error}")
                    
                    print(f"DEBUG: All WMF conversion methods failed. Skipping image.")
                    return None, None
            
            # For other formats, use standard PIL conversion
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if needed
            if image.mode in ('RGBA', 'LA', 'P'):
                image = image.convert('RGB')
            
            # Convert to bytes in the target format
            output_buffer = io.BytesIO()
            if to_mime_type == 'image/png':
                image.save(output_buffer, format='PNG')
            elif to_mime_type == 'image/jpeg':
                image.save(output_buffer, format='JPEG', quality=85)
            else:
                image.save(output_buffer, format='PNG')  # Default to PNG
            
            converted_data = output_buffer.getvalue()
            print(f"DEBUG: Successfully converted {from_mime_type} to {to_mime_type}")
            return converted_data, to_mime_type
            
        except ImportError:
            print("DEBUG: PIL not available, skipping image conversion")
            return None, None
        except Exception as e:
            print(f"DEBUG: Image conversion failed: {e}")
            return None, None
            
    except Exception as e:
        print(f"DEBUG: Error in image conversion: {e}")
        return None, None

def is_valid_image_for_gemini(mime_type, image_data):
    """Check if image is valid for Gemini API and convert if needed"""
    # Gemini supports: image/jpeg, image/png, image/webp, image/heic, image/heif
    supported_types = ['image/jpeg', 'image/png', 'image/webp', 'image/heic', 'image/heif']
    is_supported = mime_type.lower() in supported_types
    
    if not is_supported:
        print(f"DEBUG: Image rejected - unsupported MIME type: {mime_type}")
        # Try to convert unsupported formats
        if mime_type.lower() in ['image/x-wmf', 'image/wmf', 'image/bmp', 'image/tiff']:
            print(f"DEBUG: Attempting to convert {mime_type} to PNG...")
            converted_data, converted_mime = convert_image_format(image_data, mime_type, 'image/png')
            if converted_data:
                print(f"DEBUG: Successfully converted {mime_type} to PNG")
                return True  # Return True since we now have a supported format
            else:
                print(f"DEBUG: Conversion failed for {mime_type}")
                return False
    else:
        print(f"DEBUG: Image accepted - supported MIME type: {mime_type}")
    
    return is_supported

def process_images_for_gemini(images):
    """Process images to ensure they are compatible with Gemini API using multiple methods"""
    print(f"🔧 Processing {len(images)} images with enhanced methods...")
    
    # Use enhanced processing if available
    if sum(ENHANCED_IMAGE_LIBRARIES.values()) > 1:
        print("DEBUG: Using enhanced image processing with multiple libraries")
        return process_images_with_multiple_methods(images)
    else:
        print("DEBUG: Using basic image processing")
        return process_images_for_gemini_basic(images)

def process_images_for_gemini_basic(images):
    """Basic image processing (original method)"""
    valid_images = []
    for i, img in enumerate(images):
        if is_valid_image_for_gemini(img["mime_type"], img["data"]):
            valid_images.append(img)
            print(f"DEBUG: Image {i+1} is valid for Gemini API: {img['mime_type']}")
        else:
            print(f"DEBUG: Image {i+1} skipped - unsupported format: {img['mime_type']}")
    return valid_images

# --- File Extraction Agent ---

def agent_extract_content(file_stream, filename):
    """Extract text content and images from uploaded files with enhanced image processing"""
    print("AGENT [Extractor]: Starting enhanced content extraction...")
    
    try:
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'docx':
            # Handle DOCX files with image extraction
            try:
                doc = DocxDocument(file_stream)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                images = []
                
                # Extract images from DOCX
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            mime_type = rel.target_part.content_type
                            
                            if is_valid_image_for_gemini(mime_type, image_data):
                                images.append({
                                    'mime_type': mime_type,
                                    'data': base64.b64encode(image_data).decode('utf-8')
                                })
                                print(f"AGENT [Extractor]: Extracted image with MIME type: {mime_type}")
                        except Exception as e:
                            print(f"AGENT [Extractor]: Failed to extract image: {e}")
                
                print(f"AGENT [Extractor]: Extracted text ({len(text_content)} chars) and {len(images)} images from DOCX.")
                
                # Log token consumption for content extraction
                estimated_tokens = len(text_content.split()) * 1.3
                log_token_consumption("content_extraction", estimated_tokens, "local_processing", {
                    "file_type": "docx",
                    "text_length": len(text_content),
                    "images_count": len(images)
                })
                
                return text_content, images, None
            except PackageNotFoundError:
                return "", [], "Invalid DOCX file format"
            except Exception as e:
                return "", [], f"Error processing DOCX file: {str(e)}"
        
        elif file_extension == 'pdf':
            # Handle PDF files with image extraction
            try:
                pdf_reader = PyPDF2.PdfReader(file_stream)
                text_content = ""
                images = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += page.extract_text() + "\n"
                    
                    # Extract images from PDF (basic implementation)
                    # Note: PyPDF2 has limited image extraction capabilities
                    # For production, consider using pdfplumber or pdf2image
                    try:
                        if '/XObject' in page['/Resources']:
                            xObject = page['/Resources']['/XObject'].get_object()
                            for obj in xObject:
                                if xObject[obj]['/Subtype'] == '/Image':
                                    try:
                                        image_data = xObject[obj].get_data()
                                        # Determine MIME type based on image format
                                        if xObject[obj]['/Filter'] == '/DCTDecode':
                                            mime_type = 'image/jpeg'
                                        elif xObject[obj]['/Filter'] == '/FlateDecode':
                                            mime_type = 'image/png'
                                        else:
                                            mime_type = 'image/jpeg'  # Default
                                        
                                        if is_valid_image_for_gemini(mime_type, image_data):
                                            images.append({
                                                'mime_type': mime_type,
                                                'data': base64.b64encode(image_data).decode('utf-8')
                                            })
                                            print(f"AGENT [Extractor]: Extracted image from PDF page {page_num + 1}")
                                    except Exception as e:
                                        print(f"AGENT [Extractor]: Failed to extract image from PDF: {e}")
                    except Exception as e:
                        print(f"AGENT [Extractor]: Error processing PDF images: {e}")
                
                print(f"AGENT [Extractor]: Extracted text ({len(text_content)} chars) and {len(images)} images from PDF.")
                
                # Log token consumption for content extraction
                estimated_tokens = len(text_content.split()) * 1.3
                log_token_consumption("content_extraction", estimated_tokens, "local_processing", {
                    "file_type": "pdf",
                    "text_length": len(text_content),
                    "images_count": len(images)
                })
                
                return text_content, images, None
            except Exception as e:
                return "", [], f"Error processing PDF file: {str(e)}"
        
        else:
            return "", [], f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        print(f"AGENT [Extractor]: Error during extraction: {e}")
        return "", [], f"Extraction failed: {str(e)}"

# --- Gemini API Agent Caller ---

def call_generative_agent(prompt_parts, is_json=False, stage_name="generative_agent"):
    """Call the generative AI agent with proper error handling, detailed token tracking, and retry logic"""
    import time
    import random
    
    max_retries = 3
    base_delay = 2  # Base delay in seconds
    
    for attempt in range(max_retries):
        try:
            # Extract text from prompt_parts (which are dictionaries)
            text_parts = []
            total_input_chars = 0
            image_count = 0
            
            for part in prompt_parts:
                if isinstance(part, dict):
                    if 'text' in part:
                        text_parts.append(part['text'])
                        total_input_chars += len(part['text'])
                    elif 'inline_data' in part:
                        # This is an image
                        image_count += 1
                        text_parts.append("[IMAGE]")  # Placeholder for image
                        total_input_chars += 1000  # Rough estimate for image
                else:
                    text_parts.append(str(part))
                    total_input_chars += len(str(part))
            
            total_text = " ".join(text_parts)
            estimated_input_tokens = len(total_text.split()) * 1.3  # Rough token estimation
            
            print(f"🔍 {stage_name.upper()}: Starting API call (attempt {attempt + 1}/{max_retries})")
            print(f"   📝 Input characters: {total_input_chars:,}")
            print(f"   📊 Estimated input tokens: {estimated_input_tokens:,.0f}")
            print(f"   🖼️ Images in prompt: {image_count}")
            print(f"   🔑 API Key: {GEMINI_API_KEY[:10]}...")
            print(f"   🌐 API URL: {GEMINI_API_URL}")
            
            headers = {
                "Content-Type": "application/json",
                "x-goog-api-key": GEMINI_API_KEY
            }
            
            # Prepare the request payload
            if is_json:
                payload = {
                    "contents": [{
                        "role": "user",
                        "parts": [{"text": "You are a helpful AI assistant. Please respond with valid JSON only, no additional text."}]
                    }, {
                        "role": "user",
                        "parts": prompt_parts
                    }],
                    "generationConfig": {
                        "temperature": 0.1,
                        "topK": 1,
                        "topP": 0.1,
                        "maxOutputTokens": 8192,
                        "responseMimeType": "application/json"
                    }
                }
            else:
                payload = {
                    "contents": [{
                        "role": "user",
                        "parts": prompt_parts
                    }],
                    "generationConfig": {
                        "temperature": 0.7,
                        "topK": 40,
                        "topP": 0.95,
                        "maxOutputTokens": 8192
                    }
                }
            
            print(f"   📦 Payload structure: {len(payload['contents'])} content items")
            print(f"   📦 First part type: {type(prompt_parts[0]) if prompt_parts else 'None'}")
            
            response = requests.post(GEMINI_API_URL, headers=headers, json=payload, timeout=120)
            
            print(f"   📡 Response status: {response.status_code}")
            print(f"   📡 Response headers: {dict(response.headers)}")
            
            if response.status_code == 200:
                result = response.json()
                print(f"   ✅ API call successful")
                
                # Extract detailed token information
                usage_metadata = result.get('usageMetadata', {})
                input_token_count = usage_metadata.get('promptTokenCount', estimated_input_tokens)
                output_token_count = usage_metadata.get('candidatesTokenCount', 0)
                total_token_count = usage_metadata.get('totalTokenCount', input_token_count + output_token_count)
                
                # Get response text for output token estimation if not provided
                if output_token_count == 0:
                    content = result.get("candidates", [{}])[0].get("content", {})
                    parts = content.get("parts", [])
                    response_text = parts[0].get("text", "") if parts else ""
                    output_token_count = len(response_text.split()) * 1.3
                
                # Log detailed token consumption
                details = {
                    "input_tokens": input_token_count,
                    "output_tokens": output_token_count,
                    "input_characters": total_input_chars,
                    "response_length": len(response_text) if 'response_text' in locals() else 0,
                    "is_json": is_json,
                    "attempt": attempt + 1,
                    "images_in_prompt": image_count
                }
                
                log_token_consumption(stage_name, total_token_count, "gemini-1.5-pro", details)
                
                if is_json:
                    # Extract JSON from the response
                    content = result.get("candidates", [{}])[0].get("content", {})
                    parts = content.get("parts", [])
                    if parts:
                        text_content = parts[0].get("text", "")
                        print(f"AGENT [{stage_name}]: Raw response length: {len(text_content)}")
                        print(f"AGENT [{stage_name}]: Raw response preview: {text_content[:200]}...")
                        
                        # Try to extract JSON from the response
                        try:
                            # Find JSON in the response - look for the complete JSON structure
                            json_start = text_content.find('{')
                            if json_start != -1:
                                # Find the matching closing brace by counting braces
                                brace_count = 0
                                json_end = json_start
                                for i in range(json_start, len(text_content)):
                                    if text_content[i] == '{':
                                        brace_count += 1
                                    elif text_content[i] == '}':
                                        brace_count -= 1
                                        if brace_count == 0:
                                            json_end = i + 1
                                            break
                                
                                if json_end > json_start:
                                    json_str = text_content[json_start:json_end]
                                    print(f"AGENT [{stage_name}]: Extracted JSON length: {len(json_str)}")
                                    parsed_json = json.loads(json_str)
                                    print(f"AGENT [{stage_name}]: JSON parsing successful")
                                    return parsed_json
                                else:
                                    print(f"AGENT [{stage_name}]: Could not find complete JSON structure")
                                    return None
                            else:
                                print(f"AGENT [{stage_name}]: No JSON structure found in response")
                                return None
                        except json.JSONDecodeError as e:
                            print(f"AGENT [{stage_name}]: JSON parsing failed: {e}")
                            print(f"AGENT [{stage_name}]: Attempted to parse: {text_content[json_start:json_start+500]}...")
                            return None
                    else:
                        print(f"AGENT [{stage_name}]: No content parts found in response")
                        return None
                else:
                    content = result.get("candidates", [{}])[0].get("content", {})
                    parts = content.get("parts", [])
                    return parts[0].get("text", "") if parts else ""
                    
            elif response.status_code == 429:
                # Rate limit exceeded - implement exponential backoff
                error_data = response.json()
                retry_delay = error_data.get('error', {}).get('details', [{}])[0].get('retryDelay', '40s')
                
                # Parse retry delay (e.g., "40s" -> 40 seconds)
                if isinstance(retry_delay, str) and retry_delay.endswith('s'):
                    try:
                        retry_seconds = int(retry_delay[:-1])
                    except ValueError:
                        retry_seconds = base_delay * (2 ** attempt)
                else:
                    retry_seconds = base_delay * (2 ** attempt)
                
                # Add jitter to prevent thundering herd
                jitter = random.uniform(0.5, 1.5)
                actual_delay = retry_seconds * jitter
                
                print(f"⚠️ Rate limit exceeded. Retrying in {actual_delay:.1f} seconds...")
                log_token_consumption(stage_name, 0, "gemini-1.5-pro", {
                    "error": f"Rate limit exceeded (attempt {attempt + 1})",
                    "retry_delay": actual_delay
                })
                
                if attempt < max_retries - 1:
                    time.sleep(actual_delay)
                    continue
                else:
                    print(f"❌ Max retries exceeded for {stage_name}")
                    log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": "Max retries exceeded"})
                    return None
                    
            elif response.status_code == 503:
                # Service unavailable - implement longer backoff
                print(f"⚠️ Service unavailable (503). Retrying in {base_delay * (2 ** attempt):.1f} seconds...")
                log_token_consumption(stage_name, 0, "gemini-1.5-pro", {
                    "error": f"Service unavailable (503) (attempt {attempt + 1})",
                    "retry_delay": base_delay * (2 ** attempt)
                })
                
                if attempt < max_retries - 1:
                    time.sleep(base_delay * (2 ** attempt))
                    continue
                else:
                    print(f"❌ Max retries exceeded for {stage_name} (503 errors)")
                    log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": "Max retries exceeded (503)"})
                    return None
                    
            else:
                print(f"❌ API Error: {response.status_code} - {response.text}")
                log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": f"API Error {response.status_code}"})
                return None
                
        except Exception as e:
            print(f"❌ Error calling generative agent (attempt {attempt + 1}): {e}")
            import traceback
            traceback.print_exc()
            log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": str(e), "attempt": attempt + 1})
            
            if attempt < max_retries - 1:
                delay = base_delay * (2 ** attempt) + random.uniform(0, 1)
                print(f"⚠️ Retrying in {delay:.1f} seconds...")
                time.sleep(delay)
                continue
            else:
                print(f"❌ Max retries exceeded for {stage_name}")
                return None
    
    return None

# --- Specialized Generative Agents ---

def agent_planner(text_content, images):
    print("AGENT [Planner]: Starting enhanced analysis to create a comprehensive plan...")
    prompt_text = f"""You are an expert Business Analyst with 15+ years of experience in enterprise software development. 
    
Analyze the following business requirements document thoroughly, including its text and any images. Create a comprehensive, detailed plan that covers:

1. **System Overview**: High-level architecture and main components
2. **User Roles & Permissions**: All user types and their access levels
3. **Core Features**: Primary functionality and user workflows
4. **Technical Requirements**: Non-functional requirements (performance, security, scalability)
5. **Integration Points**: External systems and APIs
6. **Data Requirements**: Data models, storage, and migration needs
7. **Security Considerations**: Authentication, authorization, and compliance
8. **Deployment Strategy**: Infrastructure and deployment approach

Focus on creating a plan that is:
- Comprehensive yet concise
- Actionable for development teams
- Aligned with industry best practices
- Scalable for future enhancements

--- DOCUMENT TEXT ---
{text_content}
--- DOCUMENT IMAGES ---
"""
    prompt_parts = [{"text": prompt_text}]
    
    # Process images to ensure they are compatible with Gemini API
    valid_images = process_images_for_gemini(images)
    
    if valid_images:
        print(f"AGENT [Planner]: Including {len(valid_images)} valid images in the analysis")
        for img in valid_images:
            prompt_parts.append({
                "inline_data": {
                    "mime_type": img["mime_type"],
                    "data": img["data"]
                }
            })
    else:
        print("AGENT [Planner]: No valid images found, proceeding with text-only analysis")
        # Update prompt to indicate no images were included
        prompt_parts[0]["text"] = prompt_text.replace("including its text and any images", "including its text (no valid images were found)")
    
    plan = call_generative_agent(prompt_parts, stage_name="planning_analysis")
    if plan is None:
        # Fallback: Create a basic plan from the text content
        print("AGENT [Planner]: API failed, creating fallback plan...")
        fallback_plan = f"""# FALLBACK ANALYSIS PLAN

## System Overview
Based on the provided requirements document, this system appears to be a business application requiring analysis.

## Key Components Identified
- User interface requirements
- Data processing needs
- System integration requirements

## Technical Considerations
- Performance requirements
- Security implementation
- Scalability planning

## Next Steps
- Detailed requirements analysis needed
- Technical architecture design required
- Implementation planning necessary

Note: This is a fallback plan due to API unavailability. For comprehensive analysis, please try again when the service is available.

--- ORIGINAL TEXT ---
{text_content[:1000]}...
"""
        error = None
        print("AGENT [Planner]: Fallback plan created successfully.")
        return fallback_plan, error
    else:
        error = None
        print("AGENT [Planner]: Enhanced plan created successfully.")
        return plan, error

def agent_trd_writer(plan, original_text):
    """Enhanced TRD writer with better structure and content"""
    print("AGENT [TRD Writer]: Creating comprehensive Technical Requirements Document...")
    prompt = f"""You are a Senior Technical Writer and Business Analyst with 20+ years of experience in enterprise software documentation.

Create a comprehensive Technical Requirements Document (TRD) based on the following plan and original requirements. Structure it as follows:

# TECHNICAL REQUIREMENTS DOCUMENT

## 1. EXECUTIVE SUMMARY
- Project Overview
- Key Objectives
- Success Criteria

## 2. SYSTEM OVERVIEW
- High-Level Architecture
- System Components
- Technology Stack

## 3. FUNCTIONAL REQUIREMENTS
- User Stories and Use Cases
- Business Rules
- Data Requirements

## 4. NON-FUNCTIONAL REQUIREMENTS
- Performance Requirements
- Security Requirements
- Scalability Requirements
- Usability Requirements

## 5. TECHNICAL SPECIFICATIONS
- API Specifications
- Database Design
- Integration Points

## 6. SECURITY CONSIDERATIONS
- Authentication & Authorization
- Data Protection
- Compliance Requirements

## 7. DEPLOYMENT & OPERATIONS
- Infrastructure Requirements
- Deployment Strategy
- Monitoring & Maintenance

## 8. TESTING STRATEGY
- Test Types
- Quality Assurance
- Acceptance Criteria

Make the document:
- Comprehensive yet concise
- Technical but accessible
- Actionable for development teams
- Compliant with industry standards

--- HIGH-LEVEL PLAN ---
{plan}

--- ORIGINAL REQUIREMENTS TEXT ---
{original_text}
"""
    result = call_generative_agent([{"text": prompt}], stage_name="trd_writer")
    if result is None:
        return "TRD could not be completed due to API error.", "API call failed"
    return result, None

def agent_diagrammer(plan, diagram_type):
    """Enhanced diagram generator with better Mermaid syntax and robust error handling"""
    print(f"AGENT [Diagrammer]: Creating {diagram_type} diagram...")
    
    try:
        if diagram_type == "HLD":
            prompt = f"""You are a Senior Solution Architect specializing in system design diagrams.

Create a High-Level Design (HLD) diagram using Mermaid syntax. The diagram should show:

**System Architecture Components:**
- User Interface Layer (Web/Mobile/Desktop)
- Application Layer (API Gateway, Services)
- Business Logic Layer (Core Services)
- Data Access Layer (Repositories)
- External Systems Integration
- Security Layer (Authentication/Authorization)

**Key Elements:**
- System boundaries and data flow
- Integration points with external systems
- Technology stack components
- Security layers and authentication
- Load balancers and caching

Use Mermaid flowchart syntax with:
- Clear component names in boxes
- Proper relationships with arrows
- Color coding for different layers
- Professional styling and layout
- AVOID subgraph syntax - use simple flowchart instead

IMPORTANT: 
- Respond ONLY with the Mermaid diagram code, no additional text or explanations
- Use simple flowchart syntax, NOT subgraph
- Use clear, simple node labels
- Avoid complex syntax that might cause parsing errors

Example format:
```mermaid
flowchart TD
    A[User Interface] --> B[API Gateway]
    B --> C[Authentication Service]
    C --> D[Business Logic]
    D --> E[Data Access Layer]
    E --> F[(Database)]
    D --> G[External APIs]
```

--- HIGH-LEVEL PLAN ---
{plan}
"""
        else:  # LLD
            prompt = f"""You are a Senior Software Architect specializing in detailed system design.

Create a Low-Level Design (LLD) diagram using Mermaid syntax. The diagram should show:

**Detailed System Components:**
- Specific modules, classes, and interfaces
- Database schema relationships and tables
- API endpoints and HTTP methods
- Data flow between components
- Error handling and validation paths
- Caching and performance optimizations

**Key Elements:**
- Detailed component interactions
- Database tables and relationships
- API specifications with methods
- Security implementation details
- Performance considerations
- Error handling flows

Use Mermaid flowchart syntax with:
- Detailed component breakdown
- Clear data flow with labels
- Error handling paths
- Performance indicators
- Database relationships
- AVOID subgraph syntax - use simple flowchart instead

IMPORTANT: 
- Respond ONLY with the Mermaid diagram code, no additional text or explanations
- Use simple flowchart syntax, NOT subgraph
- Use clear, simple node labels
- Avoid complex syntax that might cause parsing errors

Example format:
```mermaid
flowchart TD
    A[UserController] --> B[UserService]
    B --> C[UserRepository]
    C --> D[(User Table)]
    B --> E[EmailService]
    B --> F[ValidationService]
    F --> G[ErrorHandler]
```

--- HIGH-LEVEL PLAN ---
{plan}
"""
        
        print(f"AGENT [Diagrammer]: Sending {diagram_type} diagram request...")
        result = call_generative_agent([{"text": prompt}], stage_name=f"{diagram_type.lower()}_diagram")
        
        if result is None:
            print(f"AGENT [Diagrammer]: {diagram_type} diagram generation failed - API returned None")
            return f"Failed to generate {diagram_type} diagram due to API error.", "API call failed"
        
        # Clean the result to extract Mermaid code
        cleaned_result = extract_mermaid_code(result)
        
        if not cleaned_result:
            print(f"AGENT [Diagrammer]: No valid Mermaid code found in {diagram_type} result")
            return f"Failed to extract valid Mermaid code from {diagram_type} response.", "Invalid Mermaid code"
        
        print(f"AGENT [Diagrammer]: Successfully generated {diagram_type} diagram")
        return cleaned_result, None
        
    except Exception as e:
        print(f"AGENT [Diagrammer]: Exception during {diagram_type} generation: {e}")
        import traceback
        traceback.print_exc()
        return f"Failed to generate {diagram_type} diagram due to exception: {str(e)}", str(e)

def agent_backlog_creator(plan, original_text, trd):
    """Enhanced backlog creator with better structure and fallback mechanism"""
    print("AGENT [Backlog Creator]: Creating comprehensive project backlog...")
    prompt = f"""You are a Senior Product Manager and Agile Coach with 15+ years of experience in software development.

Create a comprehensive project backlog based on the following plan, requirements, and technical requirements document. Structure it as a JSON object with the following format:

{{
  "backlog": [
    {{
      "id": "epic-1",
      "type": "Epic",
      "title": "Epic Title",
      "description": "Epic description",
      "priority": "High/Medium/Low",
      "effort": "Story points estimate",
      "trd_sections": ["Section 1", "Section 2"],
      "requirements_covered": ["Requirement 1", "Requirement 2"],
      "children": [
        {{
          "id": "feature-1",
          "type": "Feature",
          "title": "Feature Title",
          "description": "Feature description",
          "priority": "High/Medium/Low",
          "effort": "Story points estimate",
          "trd_sections": ["Section 1"],
          "requirements_covered": ["Requirement 1"],
          "children": [
            {{
              "id": "story-1",
              "type": "User Story",
              "title": "User Story Title",
              "description": "As a [user], I want [feature] so that [benefit]",
              "priority": "High/Medium/Low",
              "effort": "Story points estimate",
              "acceptance_criteria": ["Criterion 1", "Criterion 2"],
              "trd_sections": ["Section 1"],
              "requirements_covered": ["Requirement 1"]
            }}
          ]
        }}
      ]
    }}
  ]
}}

**Backlog Structure:**
1. **Epics**: High-level business capabilities (3-5 epics recommended)
2. **Features**: Major functionality within epics (2-4 features per epic)
3. **User Stories**: Detailed requirements with acceptance criteria (3-8 stories per feature)

**Linking Requirements:**
- For each backlog item, identify which TRD sections it addresses
- For each backlog item, identify which specific requirements it covers
- Use the TRD section headers (e.g., "FUNCTIONAL REQUIREMENTS", "SECURITY CONSIDERATIONS")
- Reference specific requirements from the original text

**Considerations:**
- Prioritize by business value and technical dependencies
- Include technical debt items
- Consider security, performance, and scalability requirements
- Provide realistic effort estimates (1-13 story points)
- Include acceptance criteria for all user stories
- Cover all major system components from the TRD
- Include testing, deployment, and maintenance stories
- Ensure each item links to specific TRD sections and requirements

**Requirements:**
- Create at least 3-5 epics covering different aspects of the system
- Each epic should have 2-4 features
- Each feature should have 3-8 user stories
- Total backlog should have 20-50 user stories
- Include stories for non-functional requirements (security, performance, etc.)
- Every item must have trd_sections and requirements_covered arrays

IMPORTANT: Respond ONLY with valid JSON, no additional text or explanations.

--- HIGH-LEVEL PLAN ---
{plan}

--- TECHNICAL REQUIREMENTS DOCUMENT ---
{trd if trd else "Technical Requirements Document not available"}

--- ORIGINAL REQUIREMENTS TEXT ---
{original_text}
"""
    result = call_generative_agent([{"text": prompt}], is_json=True, stage_name="backlog_creation")
    if result is None:
        print("AGENT [Backlog Creator]: API failed, creating fallback backlog...")
        # Create a fallback backlog structure with linking information
        fallback_backlog = {
            "backlog": [
                {
                    "id": "epic-1",
                    "type": "Epic",
                    "title": "System Implementation",
                    "description": "Core system implementation based on requirements analysis",
                    "priority": "High",
                    "effort": "40",
                    "trd_sections": ["SYSTEM OVERVIEW", "FUNCTIONAL REQUIREMENTS"],
                    "requirements_covered": ["System architecture setup", "Core functionality implementation"],
                    "children": [
                        {
                            "id": "feature-1",
                            "type": "Feature",
                            "title": "Basic System Setup",
                            "description": "Initial system setup and configuration",
                            "priority": "High",
                            "effort": "13",
                            "trd_sections": ["SYSTEM OVERVIEW"],
                            "requirements_covered": ["System architecture setup"],
                            "children": [
                                {
                                    "id": "story-1",
                                    "type": "User Story",
                                    "title": "System Architecture Setup",
                                    "description": "As a developer, I want to set up the system architecture so that the foundation is ready for development",
                                    "priority": "High",
                                    "effort": "5",
                                    "acceptance_criteria": ["Architecture diagram is created", "Technology stack is defined"],
                                    "trd_sections": ["SYSTEM OVERVIEW"],
                                    "requirements_covered": ["System architecture setup"]
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        return json.dumps(fallback_backlog), None
    
    # Debug: Log the AI result
    print(f"AGENT [Backlog Creator]: AI returned result type: {type(result)}")
    print(f"AGENT [Backlog Creator]: AI result keys: {list(result.keys()) if isinstance(result, dict) else 'Not a dict'}")
    if isinstance(result, dict) and 'backlog' in result:
        print(f"AGENT [Backlog Creator]: AI backlog has {len(result['backlog'])} items")
        if result['backlog']:
            first_item = result['backlog'][0]
            print(f"AGENT [Backlog Creator]: First item title: {first_item.get('title', 'No title')}")
            print(f"AGENT [Backlog Creator]: First item type: {first_item.get('type', 'No type')}")
    
    return result, None

# --- Helper functions (add_ids_to_backlog, Azure DevOps, etc.) ---
def add_ids_to_backlog(items):
    if not isinstance(items, list): return []
    for i, epic in enumerate(items):
        epic['id'] = f"E-{i+1}"
        for j, feature in enumerate(epic.get('children', []), 1):
            feature['id'] = f"F-{j}"
            for k, story in enumerate(feature.get('children', []), 1):
                story['id'] = f"US-{k}"
    return items

# --- Enhanced Azure DevOps Integration ---

def get_ado_headers():
    """Get Azure DevOps headers with enhanced error handling"""
    if not ADO_PAT:
        raise ValueError("Azure DevOps Personal Access Token not configured")
    return {
        'Authorization': f'Basic {base64.b64encode(f":{ADO_PAT}".encode()).decode()}',
        'Content-Type': 'application/json-patch+json'
    }

def create_ado_work_item(item_type, title, description="", parent_url=None, tags=None):
    """Enhanced Azure DevOps work item creation with better error handling"""
    try:
        headers = get_ado_headers()
        
        # Prepare work item data
        work_item_data = [
            {
                "op": "add",
                "path": "/fields/System.Title",
                "value": title
            },
            {
                "op": "add",
                "path": "/fields/System.Description",
                "value": description
            }
        ]
        
        # Add tags if provided
        if tags:
            work_item_data.append({
                "op": "add",
                "path": "/fields/System.Tags",
                "value": "; ".join(tags)
            })
        
        # Add parent relationship if provided
        if parent_url:
            work_item_data.append({
                "op": "add",
                "path": "/relations/-",
                "value": {
                    "rel": "System.LinkTypes.Hierarchy-Reverse",
                    "url": parent_url
                }
            })
        
        url = f"{ADO_ORGANIZATION_URL}{ADO_PROJECT_NAME}/_apis/wit/workitems/${item_type}?api-version=6.0"
        
        response = requests.post(url, headers=headers, json=work_item_data)
        response.raise_for_status()
        
        result = response.json()
        print(f"Successfully created {item_type} work item: {result.get('id')}")
        return result
        
    except requests.exceptions.RequestException as e:
        print(f"Error creating Azure DevOps work item: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None
    except Exception as e:
        print(f"Unexpected error in Azure DevOps integration: {e}")
        return None

def process_backlog_for_ado(backlog):
    """Enhanced backlog processing for Azure DevOps with better structure"""
    if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
        print("Azure DevOps not configured, skipping work item creation")
        return {"success": False, "message": "Azure DevOps not configured"}
    
    try:
        created_items = []
        
        for epic in backlog:
            # Create Epic
            epic_description = f"Epic: {epic.get('title', 'Untitled Epic')}\n\nGenerated by BA Agent"
            epic_result = create_ado_work_item(
                "Epic", 
                epic.get('title', 'Untitled Epic'),
                epic_description,
                tags=["BA-Agent-Generated"]
            )
            
            if epic_result:
                epic_url = epic_result.get('url')
                created_items.append({
                    "type": "Epic",
                    "id": epic_result.get('id'),
                    "title": epic.get('title'),
                    "url": epic_url
                })
                
                # Create Features under Epic
                for feature in epic.get('children', []):
                    feature_description = f"Feature: {feature.get('title', 'Untitled Feature')}\n\nParent Epic: {epic.get('title')}"
                    feature_result = create_ado_work_item(
                        "Feature",
                        feature.get('title', 'Untitled Feature'),
                        feature_description,
                        parent_url=epic_url,
                        tags=["BA-Agent-Generated"]
                    )
                    
                    if feature_result:
                        feature_url = feature_result.get('url')
                        created_items.append({
                            "type": "Feature",
                            "id": feature_result.get('id'),
                            "title": feature.get('title'),
                            "url": feature_url
                        })
                        
                        # Create User Stories under Feature
                        for story in feature.get('children', []):
                            story_description = f"User Story: {story.get('title', 'Untitled Story')}\n\nParent Feature: {feature.get('title')}"
                            story_result = create_ado_work_item(
                                "User Story",
                                story.get('title', 'Untitled Story'),
                                story_description,
                                parent_url=feature_url,
                                tags=["BA-Agent-Generated"]
                            )
                            
                            if story_result:
                                created_items.append({
                                    "type": "User Story",
                                    "id": story_result.get('id'),
                                    "title": story.get('title'),
                                    "url": story_result.get('url')
                                })
        
        return {
            "success": True,
            "message": f"Successfully created {len(created_items)} work items in Azure DevOps",
            "items": created_items
        }
        
    except Exception as e:
        print(f"Error processing backlog for Azure DevOps: {e}")
        return {"success": False, "message": f"Error: {str(e)}"}

# --- Enhanced Email Notification System ---

def send_email_notification(subject, content, recipient_email=None, attachment_data=None):
    """Enhanced email notification system with better error handling"""
    try:
        if not ACS_CONNECTION_STRING or not ACS_SENDER_ADDRESS:
            print("Email service not configured, skipping notification")
            return {"success": False, "message": "Email service not configured"}
        
        # Use provided recipient or default
        recipient = recipient_email or APPROVAL_RECIPIENT_EMAIL
        if not recipient:
            print("No recipient email configured")
            return {"success": False, "message": "No recipient email configured"}
        
        # Create email client
        email_client = EmailClient.from_connection_string(ACS_CONNECTION_STRING)
        
        # Prepare email content
        email_content = {
            "subject": subject,
            "plainText": content,
            "html": content.replace('\n', '<br>')
        }
        
        # Create message
        message = {
            "senderAddress": ACS_SENDER_ADDRESS,
            "recipients": {
                "to": [{"address": recipient}]
            },
            "content": email_content
        }
        
        # Add attachment if provided
        if attachment_data:
            message["attachments"] = [attachment_data]
        
        # Send email
        poller = email_client.begin_send(message)
        result = poller.result()
        
        print(f"Email notification sent successfully to {recipient}")
        return {"success": True, "message": "Email sent successfully"}
        
    except Exception as e:
        print(f"Error sending email notification: {e}")
        return {"success": False, "message": f"Error: {str(e)}"}

def send_analysis_completion_notification(analysis_id, filename, results_summary):
    """Send notification when analysis is completed"""
    subject = f"BA Agent Analysis Completed - {filename}"
    content = f"""
Analysis completed successfully!

File: {filename}
Analysis ID: {analysis_id}

Summary:
{results_summary}

You can view the full results in the BA Agent application.
"""
    
    return send_email_notification(subject, content)

def send_approval_notification(analysis_id, filename, approval_url):
    """Send notification for approval requests"""
    subject = f"BA Agent Analysis Ready for Approval - {filename}"
    content = f"""
A new analysis is ready for your approval.

File: {filename}
Analysis ID: {analysis_id}

Please review and approve the analysis at:
{approval_url}

This analysis includes:
- Technical Requirements Document
- High and Low Level Design Diagrams
- Project Backlog
- Risk Assessment
- Cost Estimation
- Sentiment Analysis
"""
    
    return send_email_notification(subject, content)

def extract_mermaid_code(text):
    if not text:
        return ""
    match = re.search(r"```mermaid\n([\s\S]*?)```", text)
    if match:
        return match.group(1).strip()
    # Remove any standalone ```
    text = re.sub(r"```", "", text)
    return text.strip()

def render_mermaid_to_png(mermaid_code: str) -> BytesIO:
    """Render Mermaid code to PNG with error handling and code cleaning"""
    
    def clean_mermaid_code(code):
        """Clean problematic characters from Mermaid code with better syntax preservation"""
        if not code:
            return ""
        
        cleaned = code
        
        # Step 1: Handle <br> tags properly - replace with spaces instead of newlines
        cleaned = cleaned.replace('<br>', ' ')
        cleaned = cleaned.replace('<br/>', ' ')
        cleaned = cleaned.replace('<br />', ' ')
        
        # Step 2: Handle specific problematic patterns in node labels
        # Pattern: A[ASP Pages <br> (e.g., Rlv_ISLLPOL_2] -> A[ASP Pages]
        import re
        cleaned = re.sub(r'([A-Z])\[([^\]]*?)(?:<br>|\([^)]*\))[^\]]*?\]', lambda m: m.group(1) + '[' + m.group(2).strip() + ']', cleaned)
        
        # Step 3: Handle patterns like Q[/policies (GET)] -> Q[policies]
        cleaned = re.sub(r'([A-Z])\[([^\]]*?)\/\([^)]*\)([^\]]*?)\]', lambda m: m.group(1) + '[' + (m.group(2) + m.group(3)).strip() + ']', cleaned)
        
        # Step 4: Remove any remaining problematic characters but preserve structure
        cleaned = cleaned.replace('<', '')  # Remove angle brackets
        cleaned = cleaned.replace('>', '')
        cleaned = cleaned.replace('/', '')  # Remove forward slashes
        cleaned = cleaned.replace('\\', '')  # Remove backslashes
        cleaned = ' '.join(cleaned.split())  # Normalize whitespace
        cleaned = cleaned.strip()
        
        return cleaned
    
    # Try original code first
    url = "https://kroki.io/mermaid/png"
    headers = {"Content-Type": "text/plain"}
    
    try:
        response = requests.post(url, data=mermaid_code.encode("utf-8"), headers=headers, timeout=30)
        if response.status_code == 200:
            return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to render original Mermaid code: {e}")
    
    # Try with cleaned code as fallback
    try:
        cleaned_code = clean_mermaid_code(mermaid_code)
        if cleaned_code and cleaned_code != mermaid_code:
            print(f"Attempting to render cleaned Mermaid code...")
            response = requests.post(url, data=cleaned_code.encode("utf-8"), headers=headers, timeout=30)
            if response.status_code == 200:
                return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to render cleaned Mermaid code: {e}")
    
    # If all else fails, create a placeholder image
    raise ValueError(f"Kroki diagram generation failed for both original and cleaned code")

# --- API Endpoints ---

@app.route("/api/test_image_extraction", methods=['POST'])
def test_image_extraction():
    """Test endpoint to check image extraction from documents"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        print(f"🧪 Testing image extraction from: {file.filename}")
        
        # Test the extraction function
        text_content, images, error = agent_extract_content(io.BytesIO(file.read()), file.filename)
        
        if error:
            return jsonify({
                "success": False,
                "error": error,
                "filename": file.filename
            })
        
        return jsonify({
            "success": True,
            "filename": file.filename,
            "text_length": len(text_content),
            "images_count": len(images),
            "images": [
                {
                    "mime_type": img.get('mime_type', 'unknown'),
                    "data_length": len(img.get('data', ''))
                } for img in images
            ],
            "text_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content
        })
        
    except Exception as e:
        print(f"❌ Exception during image extraction test: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        })

@app.route("/api/test_gemini", methods=['GET'])
def test_gemini_api():
    """Test endpoint to check if Gemini API is working"""
    try:
        print("🧪 Testing Gemini API connection...")
        
        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": GEMINI_API_KEY
        }
        
        payload = {
            "contents": [{
                "role": "user",
                "parts": [{"text": "Hello, please respond with 'API is working' if you can see this message."}]
            }],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 100
            }
        }
        
        print(f"🔑 Using API key: {GEMINI_API_KEY[:10]}...")
        print(f"🌐 API URL: {GEMINI_API_URL}")
        
        response = requests.post(GEMINI_API_URL, headers=headers, json=payload, timeout=30)
        
        print(f"📡 Response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            content = result.get("candidates", [{}])[0].get("content", {})
            parts = content.get("parts", [])
            response_text = parts[0].get("text", "") if parts else ""
            
            return jsonify({
                "success": True,
                "message": "API is working",
                "response": response_text,
                "status_code": response.status_code
            })
        else:
            print(f"❌ API Error: {response.status_code} - {response.text}")
            return jsonify({
                "success": False,
                "error": f"API Error {response.status_code}",
                "response_text": response.text,
                "status_code": response.status_code
            })
            
    except Exception as e:
        print(f"❌ Exception during API test: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e),
            "status_code": 500
        })

@app.route("/api/test_diagrams", methods=['POST'])
def test_diagram_generation():
    """Test endpoint to check HLD and LLD diagram generation"""
    try:
        data = request.get_json()
        diagram_type = data.get('type', 'HLD')  # HLD or LLD
        test_plan = data.get('plan', 'Create a simple web application with user authentication and database storage.')
        
        print(f"🧪 Testing {diagram_type} diagram generation...")
        
        # Test the diagrammer function
        result, error = agent_diagrammer(test_plan, diagram_type)
        
        if error:
            return jsonify({
                "success": False,
                "diagram_type": diagram_type,
                "error": error,
                "result": result
            })
        else:
            return jsonify({
                "success": True,
                "diagram_type": diagram_type,
                "result": result,
                "error": None
            })
            
    except Exception as e:
        print(f"❌ Exception during diagram test: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e),
            "diagram_type": "unknown"
        })

@app.route("/api/debug_diagrams", methods=['GET'])
def debug_diagram_generation():
    """Debug endpoint to check diagram generation status"""
    try:
        # Test with a simple plan
        test_plan = "Create a web application with user authentication, database storage, and API endpoints."
        
        print("🔍 Debugging diagram generation...")
        
        # Test HLD
        hld_result, hld_error = agent_diagrammer(test_plan, "HLD")
        
        # Test LLD
        lld_result, lld_error = agent_diagrammer(test_plan, "LLD")
        
        return jsonify({
            "hld": {
                "success": not hld_error,
                "result": hld_result,
                "error": hld_error
            },
            "lld": {
                "success": not lld_error,
                "result": lld_result,
                "error": lld_error
            },
            "test_plan": test_plan
        })
        
    except Exception as e:
        print(f"❌ Exception during diagram debug: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": str(e)
        })

@app.route("/")
def index():
    return "<h1>AI Business Analyst Backend is running!</h1>"

@app.route("/api/generate", methods=['POST'])
def orchestrator():
    print("\n--- ORCHESTRATOR: Enhanced generation task started ---")
    if 'file' not in request.files:
        print("No file part in the request")
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    text_content, images, error = agent_extract_content(io.BytesIO(file.read()), file.filename)
    if error:
        print(f"Extraction error: {error}")
        return jsonify({"error": error}), 500

    # 1. Enhanced Planning Agent
    plan, error = agent_planner(text_content, images)
    if error:
        print(f"Planning Agent failed: {error}")
        return jsonify({"error": f"Planning Agent failed: {error}"}), 500

    # 2. Specialist Agents (Enhanced with better error handling)
    print("AGENT [Orchestrator]: Starting specialist agents...")
    
    trd, err_trd = agent_trd_writer(plan, text_content)
    print(f"AGENT [Orchestrator]: TRD generation {'✅ Success' if not err_trd else f'❌ Failed: {err_trd}'}")
    
    hld, err_hld = agent_diagrammer(plan, "HLD")
    print(f"AGENT [Orchestrator]: HLD generation {'✅ Success' if not err_hld else f'❌ Failed: {err_hld}'}")
    
    lld, err_lld = agent_diagrammer(plan, "LLD")
    print(f"AGENT [Orchestrator]: LLD generation {'✅ Success' if not err_lld else f'❌ Failed: {err_lld}'}")
    
    # Pass the generated TRD to backlog creator for better backlog generation
    backlog_json, err_backlog = agent_backlog_creator(plan, text_content, trd)
    print(f"AGENT [Orchestrator]: Backlog generation {'✅ Success' if not err_backlog else f'❌ Failed: {err_backlog}'}")
    
    # Debug: Log what we received from the backlog creator
    print(f"AGENT [Orchestrator]: Backlog creator returned:")
    print(f"  Type: {type(backlog_json)}")
    print(f"  Error: {err_backlog}")
    if isinstance(backlog_json, dict):
        print(f"  Keys: {list(backlog_json.keys())}")
        if 'backlog' in backlog_json:
            print(f"  Backlog length: {len(backlog_json['backlog'])}")
    elif isinstance(backlog_json, str):
        print(f"  String length: {len(backlog_json)}")
        print(f"  First 200 chars: {backlog_json[:200]}...")

    print(f"AGENT [Orchestrator]: Agent results:")
    print(f"  TRD: {'✅ Success' if not err_trd else f'❌ Failed: {err_trd}'}")
    print(f"  HLD: {'✅ Success' if not err_hld else f'❌ Failed: {err_hld}'}")
    print(f"  LLD: {'✅ Success' if not err_lld else f'❌ Failed: {err_lld}'}")
    print(f"  Backlog: {'✅ Success' if not err_backlog else f'❌ Failed: {err_backlog}'}")

    # Count successful and failed agents
    successful_agents = sum([1 for err in [err_trd, err_hld, err_lld, err_backlog] if not err])
    total_agents = 4
    
    print(f"AGENT [Orchestrator]: {successful_agents}/{total_agents} agents completed successfully")
    
    # Continue with partial results if at least some agents succeeded
    if successful_agents == 0:
        print("AGENT [Orchestrator]: All specialist agents failed")
        return jsonify({"error": "All specialist agents failed. Please try again."}), 500
    elif successful_agents < total_agents:
        print(f"AGENT [Orchestrator]: Partial success - {successful_agents}/{total_agents} agents completed")
    
    # 3. Final Assembly with Enhanced Results
    try:
        # Handle backlog parsing with better error handling
        actual_backlog_list = []
        if not err_backlog and backlog_json:
            try:
                # Check if backlog_json is already a dictionary (from call_generative_agent with is_json=True)
                if isinstance(backlog_json, dict):
                    actual_backlog_list = backlog_json.get('backlog', [])
                    print("AGENT [Orchestrator]: Backlog is already a dictionary, extracted backlog list")
                else:
                    # Parse JSON string
                    backlog_data = json.loads(backlog_json)
                    actual_backlog_list = backlog_data.get('backlog', [])
                    print("AGENT [Orchestrator]: Parsed backlog from JSON string")
                
                if not isinstance(actual_backlog_list, list):
                    print("Warning: Backlog from AI is not a list, correcting.")
                    actual_backlog_list = []
            except json.JSONDecodeError as jde:
                print(f"Backlog Creator Agent returned invalid JSON: {backlog_json}\nError: {jde}")
                actual_backlog_list = []
        else:
            print("Backlog creation failed, using empty backlog")
        
        # Clean diagram results with better error handling
        hld_clean = ""
        lld_clean = ""
        
        if not err_hld and hld:
            hld_clean = extract_mermaid_code(hld)
            print(f"AGENT [Orchestrator]: HLD cleaned result: {'✅ Available' if hld_clean else '❌ No valid Mermaid code'}")
        else:
            print(f"AGENT [Orchestrator]: HLD failed or empty: {err_hld}")
            
        if not err_lld and lld:
            lld_clean = extract_mermaid_code(lld)
            print(f"AGENT [Orchestrator]: LLD cleaned result: {'✅ Available' if lld_clean else '❌ No valid Mermaid code'}")
        else:
            print(f"AGENT [Orchestrator]: LLD failed or empty: {err_lld}")
        
        print(f"AGENT [Orchestrator]: Final assembly:")
        print(f"  TRD: {'✅ Available' if trd else '❌ Not available'}")
        print(f"  HLD: {'✅ Available' if hld_clean else '❌ Not available'}")
        print(f"  LLD: {'✅ Available' if lld_clean else '❌ Not available'}")
        print(f"  Backlog: {'✅ Available' if actual_backlog_list else '❌ Not available'}")

        final_response = {
            "trd": trd or "Technical Requirements Document could not be generated due to API error.",
            "hld": hld_clean or "High-Level Design diagram could not be generated due to API error.",
            "lld": lld_clean or "Low-Level Design diagram could not be generated due to API error.",
            "images": images,
            "backlog": add_ids_to_backlog(actual_backlog_list),
            "enhanced_plan": plan or "",
            "partial_success": successful_agents < total_agents,
            "successful_agents": successful_agents,
            "total_agents": total_agents,
            "agent_errors": {
                "trd": err_trd,
                "hld": err_hld,
                "lld": err_lld,
                "backlog": err_backlog
            }
        }
        
        # Debug: Print the final backlog structure
        print(f"AGENT [Orchestrator]: Final backlog structure:")
        print(f"  Type: {type(final_response['backlog'])}")
        print(f"  Length: {len(final_response['backlog']) if isinstance(final_response['backlog'], list) else 'Not a list'}")
        if final_response['backlog']:
            print(f"  First item type: {type(final_response['backlog'][0])}")
            print(f"  First item keys: {list(final_response['backlog'][0].keys()) if isinstance(final_response['backlog'][0], dict) else 'Not a dict'}")
        
        # Debug: Print a sample of the backlog JSON
        try:
            backlog_json_sample = json.dumps(final_response['backlog'][:1], indent=2) if final_response['backlog'] else "[]"
            print(f"AGENT [Orchestrator]: Sample backlog JSON (first item):")
            print(backlog_json_sample[:500] + "..." if len(backlog_json_sample) > 500 else backlog_json_sample)
        except Exception as e:
            print(f"AGENT [Orchestrator]: Error serializing backlog sample: {e}")
    except Exception as e:
        import traceback
        print(f"Error during final assembly: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Error during final assembly: {e}"}), 500

    print("--- ORCHESTRATOR: All enhanced agents completed. Task finished. ---\n")
    
    # Save analysis results first
    try:
        analysis_id = save_analysis_results(final_response, text_content, file.filename if file else "requirements.txt")
        final_response["analysis_id"] = analysis_id
        print(f"AGENT [Orchestrator]: Analysis saved with ID: {analysis_id}")
    except Exception as e:
        print(f"AGENT [Orchestrator]: Warning - Failed to save analysis results: {e}")
        analysis_id = "unknown"
        final_response["analysis_id"] = analysis_id
    
    # Log total token consumption for this analysis
    try:
        total_tokens = sum(log['tokens_used'] for log in token_consumption_logs)
        log_token_consumption("analysis_complete", total_tokens, "gemini-1.5-pro", {
            "total_stages": len(token_consumption_logs),
            "analysis_id": analysis_id,
            "filename": file.filename if file else "requirements.txt"
        })
    except Exception as e:
        print(f"AGENT [Orchestrator]: Warning - Failed to log token consumption: {e}")
    
    # Send completion notification (optional - don't fail if this fails)
    try:
        results_summary = f"""
Analysis completed for {file.filename if file else 'requirements.txt'}

Generated:
- Technical Requirements Document
- High and Low Level Design Diagrams  
- Project Backlog with {len(actual_backlog_list)} epics

Analysis ID: {analysis_id}
"""
        
        email_result = send_analysis_completion_notification(
            analysis_id, 
            file.filename if file else "requirements.txt",
            results_summary
        )
        
        if email_result["success"]:
            print("Email notification sent successfully")
        else:
            print(f"Email notification failed: {email_result['message']}")
    except Exception as e:
        print(f"AGENT [Orchestrator]: Warning - Failed to send email notification: {e}")
    
    print("AGENT [Orchestrator]: Returning successful response to client")
    return jsonify(final_response)

@app.route("/api/render_mermaid", methods=["POST"])
def render_mermaid():
    data = request.get_json()
    mermaid_code = data.get("code")
    if not mermaid_code:
        return jsonify({"error": "No Mermaid code provided"}), 400
    
    try:
        png_bytes = render_mermaid_to_png(mermaid_code)
        png_bytes.seek(0)
        return send_file(png_bytes, mimetype="image/png")
    except Exception as e:
        print(f"Error rendering Mermaid diagram: {e}")
        # Return a more detailed error response
        return jsonify({
            "error": f"Failed to render diagram: {str(e)}",
            "mermaid_code_preview": mermaid_code[:200] + "..." if len(mermaid_code) > 200 else mermaid_code
        }), 500

@app.route("/api/convert_to_docx", methods=['POST'])
def convert_to_docx():
    """Convert markdown content to DOCX format"""
    try:
        data = request.get_json()
        markdown_content = data.get('markdown', '')
        
        if not markdown_content:
            return jsonify({"error": "No markdown content provided"}), 400
        
        # Convert markdown to DOCX
        doc = markdown_to_docx(markdown_content)
        
        # Save to BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return send_file(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='Technical_Requirements_Document.docx'
        )
        
    except Exception as e:
        print(f"Error converting to DOCX: {e}")
        return jsonify({"error": f"Failed to convert to DOCX: {str(e)}"}), 500

@app.route("/api/approve", methods=['POST'])
def send_for_approval():
    """Enhanced approval workflow with database storage"""
    try:
        data = request.get_json()
        analysis_id = data.get('analysis_id')
        results = data.get('results')
        
        if not analysis_id or not results:
            return jsonify({"error": "Missing analysis_id or results"}), 400
        
        # Generate unique approval ID
        approval_id = str(uuid.uuid4())
        
        # Create approval record with enhanced status tracking
        approval_record = {
            "id": approval_id,
            "analysis_id": analysis_id,
            "status": "pending",
            "created_date": datetime.now().isoformat(),  # Changed from created_at
            "updated_date": datetime.now().isoformat(),  # Changed from updated_at
            "approver_email": APPROVAL_RECIPIENT_EMAIL,
            "results_summary": {
                "has_trd": bool(results.get('trd')),
                "has_diagrams": bool(results.get('hld') or results.get('lld')),
                "has_backlog": bool(results.get('backlog')),
                "backlog_epics_count": len(results.get('backlog', []))
            }
        }
        
        # Store approval record in database
        try:
            db = next(get_db())
            save_approval_to_db(db, approval_record)
            print(f"✅ Approval saved to database: {approval_id}")
        except Exception as e:
            print(f"❌ Error saving approval to database: {e}")
            # Fallback to in-memory storage
            approval_statuses[approval_id] = approval_record
            print(f"⚠️ Using in-memory fallback for approval: {approval_id}")
        
        # Generate approval URL
        approval_url = f"{BACKEND_BASE_URL}/api/approval_response?approval_id={approval_id}"
        
        # Send approval notification email
        email_result = send_approval_notification(analysis_id, "Analysis Results", approval_url)
        
        if email_result["success"]:
            print(f"✅ Approval notification sent successfully for approval ID: {approval_id}")
        else:
            print(f"❌ Approval notification failed: {email_result['message']}")
        
        return jsonify({
            "success": True,
            "approval_id": approval_id,
            "message": "Analysis sent for approval successfully",
            "approval_url": approval_url
        })
        
    except Exception as e:
        print(f"❌ Error in approval workflow: {e}")
        return jsonify({"error": f"Failed to send for approval: {str(e)}"}), 500

@app.route("/api/approval_response")
def handle_approval_response():
    """Enhanced approval response handler with database storage"""
    try:
        approval_id = request.args.get('approval_id')
        action = request.args.get('action', 'view')  # view, approve, reject
        
        # Try to get approval from database first
        db = next(get_db())
        approval_record = get_approval_from_db(db, approval_id)
        
        # Fallback to in-memory storage if not in database
        if not approval_record and approval_id in approval_statuses:
            approval_record = approval_statuses[approval_id]
        
        if not approval_record:
            return jsonify({"error": "Invalid approval ID"}), 404
        
        if action == 'view':
            # Show approval interface
            return f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>BA Agent - Analysis Approval</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    .container {{ max-width: 800px; margin: 0 auto; }}
                    .header {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
                    .summary {{ background: #e9ecef; padding: 15px; border-radius: 8px; margin-bottom: 20px; }}
                    .actions {{ display: flex; gap: 10px; margin-top: 20px; }}
                    .btn {{ padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; text-decoration: none; }}
                    .btn-approve {{ background: #28a745; color: white; }}
                    .btn-reject {{ background: #dc3545; color: white; }}
                    .btn-view {{ background: #007bff; color: white; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h1>BA Agent Analysis Approval</h1>
                        <p><strong>Approval ID:</strong> {approval_id}</p>
                        <p><strong>Analysis ID:</strong> {approval_record['analysis_id']}</p>
                        <p><strong>Created:</strong> {approval_record['created_date']}</p>
                    </div>
                    
                    <div class="summary">
                        <h2>Analysis Summary</h2>
                        <ul>
                            <li>Technical Requirements Document: {'✓' if approval_record['results_summary']['has_trd'] else '✗'}</li>
                            <li>Design Diagrams: {'✓' if approval_record['results_summary']['has_diagrams'] else '✗'}</li>
                            <li>Project Backlog: {'✓' if approval_record['results_summary']['has_backlog'] else '✗'}</li>
                            <li>Backlog Epics: {approval_record['results_summary']['backlog_epics_count']}</li>
                        </ul>
                    </div>
                    
                    <div class="actions">
                        <a href="?approval_id={approval_id}&action=approve" class="btn btn-approve">Approve Analysis</a>
                        <a href="?approval_id={approval_id}&action=reject" class="btn btn-reject">Reject Analysis</a>
                        <a href="/" class="btn btn-view">View in Application</a>
                    </div>
                </div>
            </body>
            </html>
            """
        
        elif action in ['approve', 'reject']:
            # Update approval status
            update_data = {
                'status': action,
                'updated_date': datetime.now().isoformat(),
                'approver_response': action
            }
            
            # Try to update in database first
            db_updated = update_approval_in_db(db, approval_id, update_data)
            
            # Fallback to in-memory update
            if not db_updated:
                approval_record['status'] = action
                approval_record['updated_date'] = datetime.now().isoformat()
                approval_record['approver_response'] = action
                approval_statuses[approval_id] = approval_record
            
            # If approved, automatically create ADO work items
            ado_result = None
            if action == 'approve':
                try:
                    # Get the analysis results to extract backlog
                    analysis_id = approval_record['analysis_id']
                    analysis = get_analysis_by_id(analysis_id)
                    
                    print(f"Processing approval for analysis {analysis_id}")
                    print(f"Analysis found: {analysis is not None}")
                    
                    if analysis and analysis.get('results', {}).get('backlog'):
                        backlog = analysis['results']['backlog']
                        print(f"Backlog found with {len(backlog)} epics")
                        
                        # Check Azure DevOps configuration first
                        if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
                            print("Azure DevOps not configured - skipping work item creation")
                            ado_result = {
                                "success": False, 
                                "message": "Azure DevOps not configured. Please set ADO_PERSONAL_ACCESS_TOKEN, ADO_ORGANIZATION_URL, and ADO_PROJECT_NAME environment variables."
                            }
                        else:
                            print("Azure DevOps configured - creating work items...")
                            ado_result = process_backlog_for_ado(backlog)
                            
                            if ado_result.get('success'):
                                items_created = len(ado_result.get('items', []))
                                print(f"Successfully created {items_created} ADO work items")
                                
                                # Log details of created items
                                for item in ado_result.get('items', []):
                                    print(f"  Created {item['type']}: {item['title']} (ID: {item['id']})")
                            else:
                                print(f"ADO work item creation failed: {ado_result.get('message')}")
                    else:
                        print(f"No backlog found for analysis {analysis_id}")
                        ado_result = {
                            "success": False, 
                            "message": "No backlog found in analysis results"
                        }
                except Exception as e:
                    print(f"Error creating ADO work items: {e}")
                    import traceback
                    traceback.print_exc()
                    ado_result = {"success": False, "message": f"Error: {str(e)}"}
            
            # Update ADO result in database
            if ado_result:
                update_approval_in_db(db, approval_id, {'ado_result': ado_result})
            
            # Send notification email
            subject = f"Analysis {'Approved' if action == 'approve' else 'Rejected'} - {approval_record['analysis_id']}"
            content = f"""
Analysis has been {action}ed.

Approval ID: {approval_id}
Analysis ID: {approval_record['analysis_id']}
Status: {action.upper()}
Updated: {approval_record['updated_date']}

You can view the full results in the BA Agent application.
"""
            
            # Add ADO information to email if approved
            if action == 'approve' and ado_result:
                if ado_result.get('success'):
                    content += f"\n\nADO Work Items Created: {len(ado_result.get('items', []))} items"
                    content += f"\nADO Status: Success"
                else:
                    content += f"\n\nADO Work Items: Failed to create"
                    content += f"\nADO Error: {ado_result.get('message')}"
            
            email_result = send_email_notification(subject, content)
            
            if email_result["success"]:
                print(f"✅ Approval {action} notification sent successfully")
            else:
                print(f"❌ Approval {action} notification failed: {email_result['message']}")
            
            return jsonify({
                "success": True,
                "status": action,
                "message": f"Analysis {action}ed successfully",
                "ado_result": ado_result
            })
        
        else:
            return jsonify({"error": "Invalid action"}), 400
            
    except Exception as e:
        print(f"❌ Error in approval response handler: {e}")
        return jsonify({"error": f"Failed to process approval response: {str(e)}"}), 500

@app.route("/api/approval_status/<approval_id>")
def get_approval_status(approval_id):
    """Get approval status with database storage"""
    try:
        # Try to get approval from database first
        db = next(get_db())
        approval_record = get_approval_from_db(db, approval_id)
        
        # Fallback to in-memory storage if not in database
        if not approval_record and approval_id in approval_statuses:
            approval_record = approval_statuses[approval_id]
        
        if not approval_record:
            return jsonify({"error": "Approval not found"}), 404
        
        # If approved, try to get Azure DevOps results
        ado_results = None
        if approval_record.get('status') == 'approve':
            try:
                # Get the analysis results to extract backlog
                analysis_id = approval_record['analysis_id']
                analysis = get_analysis_by_id(analysis_id)
                
                if analysis and analysis.get('results', {}).get('backlog'):
                    # Check if we already have ADO results in the approval record
                    if 'ado_result' in approval_record:
                        ado_results = approval_record['ado_result']
                    else:
                        # Try to create ADO work items if not already done
                        if ADO_PAT and ADO_ORGANIZATION_URL and ADO_PROJECT_NAME:
                            ado_results = process_backlog_for_ado(analysis['results']['backlog'])
                            # Store the results in the approval record
                            update_approval_in_db(db, approval_id, {'ado_result': ado_results})
                        else:
                            ado_results = {
                                "success": False,
                                "message": "Azure DevOps not configured"
                            }
            except Exception as e:
                ado_results = {
                    "success": False,
                    "message": f"Error processing Azure DevOps: {str(e)}"
                }
        
        return jsonify({
            **approval_record,
            "ado_result": ado_results
        })
        
    except Exception as e:
        print(f"❌ Error getting approval status: {e}")
        return jsonify({"error": f"Failed to get approval status: {str(e)}"}), 500

@app.route("/api/search", methods=['POST'])
def semantic_search():
    """Perform semantic search across documents and analyses"""
    try:
        data = request.get_json()
        query = data.get('query', '')
        collection = data.get('collection', 'documents')
        n_results = data.get('n_results', 5)
        
        if not query:
            return jsonify({"error": "Query is required"}), 400
        
        # Perform vector search
        results = search_vector_db(query, collection, n_results)
        
        return jsonify({
            "query": query,
            "results": results,
            "total_results": len(results.get('documents', []))
        })
        
    except Exception as e:
        print(f"Error in semantic search: {e}")
        return jsonify({"error": f"Search failed: {str(e)}"}), 500

@app.route("/api/vector/collections", methods=['GET'])
def get_collections():
    """Get all Qdrant collections"""
    try:
        collections = qdrant_client.get_collections()
        return jsonify({
            "collections": [col.name for col in collections.collections]
        })
    except Exception as e:
        print(f"Error getting collections: {e}")
        return jsonify({"error": f"Failed to get collections: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/info", methods=['GET'])
def get_collection_info(collection_name):
    """Get information about a specific collection"""
    try:
        collection_info = qdrant_client.get_collection(collection_name)
        return jsonify({
            "name": collection_name,
            "vectors_count": collection_info.vectors_count,
            "points_count": collection_info.points_count,
            "status": collection_info.status
        })
    except Exception as e:
        print(f"Error getting collection info: {e}")
        return jsonify({"error": f"Failed to get collection info: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/points", methods=['GET'])
def get_collection_points(collection_name):
    """Get points from a collection"""
    try:
        limit = request.args.get('limit', 100, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        points = qdrant_client.scroll(
            collection_name=collection_name,
            limit=limit,
            offset=offset
        )
        
        return jsonify({
            "collection": collection_name,
            "points": [
                {
                    "id": point.id,
                    "payload": point.payload,
                    "score": getattr(point, 'score', None)
                }
                for point in points[0]
            ],
            "total": len(points[0])
        })
    except Exception as e:
        print(f"Error getting collection points: {e}")
        return jsonify({"error": f"Failed to get collection points: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/points/<point_id>", methods=['DELETE'])
def delete_point(collection_name, point_id):
    """Delete a point from a collection"""
    try:
        success = delete_from_vector_db(point_id, collection_name)
        if success:
            return jsonify({"message": f"Point {point_id} deleted successfully"})
        else:
            return jsonify({"error": "Failed to delete point"}), 500
    except Exception as e:
        print(f"Error deleting point: {e}")
        return jsonify({"error": f"Failed to delete point: {str(e)}"}), 500

# --- Enhanced Image Processing ---
def setup_enhanced_image_processing():
    """Setup enhanced image processing with multiple libraries"""
    print("🔧 Setting up enhanced image processing...")
    
    available_libraries = {
        'PIL/Pillow': False,
        'ImageMagick (Wand)': False,
        'OpenCV': False,
        'pdf2image': False
    }
    
    try:
        from PIL import Image, ImageDraw
        available_libraries['PIL/Pillow'] = True
        print("✅ PIL/Pillow available")
    except ImportError:
        print("❌ PIL/Pillow not available")
    
    try:
        from wand.image import Image as WandImage
        available_libraries['ImageMagick (Wand)'] = True
        print("✅ ImageMagick (Wand) available")
    except ImportError:
        print("❌ ImageMagick (Wand) not available")
    
    try:
        import cv2
        available_libraries['OpenCV'] = True
        print("✅ OpenCV available")
    except ImportError:
        print("❌ OpenCV not available")
    
    try:
        from pdf2image import convert_from_bytes
        available_libraries['pdf2image'] = True
        print("✅ pdf2image available")
    except ImportError:
        print("❌ pdf2image not available")
    
    print(f"📊 Available libraries: {sum(available_libraries.values())}/{len(available_libraries)}")
    return available_libraries

def process_images_with_multiple_methods(images):
    """Process images using multiple conversion methods"""
    processed_images = []
    
    for img in images:
        mime_type = img.get('mime_type', '')
        image_data = base64.b64decode(img.get('data', ''))
        
        # Method 1: Try PIL
        processed_img = convert_with_pil(image_data, mime_type)
        if processed_img:
            processed_images.append(processed_img)
            continue
        
        # Method 2: Try ImageMagick
        processed_img = convert_with_wand(image_data, mime_type)
        if processed_img:
            processed_images.append(processed_img)
            continue
        
        # Method 3: Try OpenCV
        processed_img = convert_with_opencv(image_data, mime_type)
        if processed_img:
            processed_images.append(processed_img)
            continue
        
        # Method 4: Create placeholder
        processed_img = create_placeholder(mime_type)
        if processed_img:
            processed_images.append(processed_img)
    
    return processed_images

def convert_with_pil(image_data, mime_type):
    """Convert image using PIL"""
    try:
        from PIL import Image
        import io
        
        image = Image.open(io.BytesIO(image_data))
        if image.mode in ('RGBA', 'LA', 'P'):
            image = image.convert('RGB')
        
        output_buffer = io.BytesIO()
        image.save(output_buffer, format='PNG')
        converted_data = output_buffer.getvalue()
        
        return {
            'mime_type': 'image/png',
            'data': base64.b64encode(converted_data).decode('utf-8')
        }
    except Exception as e:
        print(f"DEBUG: PIL conversion failed for {mime_type}: {e}")
        return None

def convert_with_wand(image_data, mime_type):
    """Convert image using ImageMagick"""
    try:
        from wand.image import Image as WandImage
        
        with WandImage(blob=image_data) as wand_img:
            wand_img.format = 'png'
            converted_data = wand_img.make_blob()
            
            return {
                'mime_type': 'image/png',
                'data': base64.b64encode(converted_data).decode('utf-8')
            }
    except ImportError:
        print("DEBUG: ImageMagick (wand) not available")
        return None
    except Exception as e:
        print(f"DEBUG: ImageMagick conversion failed for {mime_type}: {e}")
        return None

def convert_with_opencv(image_data, mime_type):
    """Convert image using OpenCV"""
    try:
        import cv2
        import numpy as np
        
        nparr = np.frombuffer(image_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return None
        
        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        success, buffer = cv2.imencode('.png', img_rgb)
        
        if success:
            converted_data = buffer.tobytes()
            return {
                'mime_type': 'image/png',
                'data': base64.b64encode(converted_data).decode('utf-8')
            }
    except ImportError:
        print("DEBUG: OpenCV not available")
        return None
    except Exception as e:
        print(f"DEBUG: OpenCV conversion failed for {mime_type}: {e}")
        return None

def create_placeholder(mime_type):
    """Create placeholder image"""
    try:
        from PIL import Image, ImageDraw
        
        placeholder = Image.new('RGB', (400, 300), color='lightgray')
        draw = ImageDraw.Draw(placeholder)
        
        draw.text((20, 20), f"Image: {mime_type}", fill='black')
        draw.text((20, 50), "Format not supported", fill='red')
        draw.text((20, 80), "Converted to placeholder", fill='blue')
        
        output_buffer = io.BytesIO()
        placeholder.save(output_buffer, format='PNG')
        converted_data = output_buffer.getvalue()
        
        return {
            'mime_type': 'image/png',
            'data': base64.b64encode(converted_data).decode('utf-8')
        }
    except Exception as e:
        print(f"DEBUG: Failed to create placeholder: {e}")
        return None

# Initialize enhanced image processing
ENHANCED_IMAGE_LIBRARIES = setup_enhanced_image_processing()

@app.route("/api/ado/status", methods=['GET'])
def check_ado_status():
    """Check Azure DevOps configuration and connection status"""
    try:
        # Check if Azure DevOps is configured
        if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
            return jsonify({
                "configured": False,
                "message": "Azure DevOps not configured",
                "missing": []
            })
        
        # Check what's missing
        missing = []
        if not ADO_PAT:
            missing.append("ADO_PERSONAL_ACCESS_TOKEN")
        if not ADO_ORGANIZATION_URL:
            missing.append("ADO_ORGANIZATION_URL")
        if not ADO_PROJECT_NAME:
            missing.append("ADO_PROJECT_NAME")
        
        if missing:
            return jsonify({
                "configured": False,
                "message": f"Missing configuration: {', '.join(missing)}",
                "missing": missing
            })
        
        # Test connection by trying to get project info
        try:
            headers = get_ado_headers()
            test_url = f"{ADO_ORGANIZATION_URL}{ADO_PROJECT_NAME}/_apis/project?api-version=6.0"
            response = requests.get(test_url, headers=headers)
            
            if response.status_code == 200:
                project_info = response.json()
                return jsonify({
                    "configured": True,
                    "connected": True,
                    "message": "Azure DevOps connection successful",
                    "project": {
                        "name": project_info.get('name'),
                        "id": project_info.get('id'),
                        "description": project_info.get('description')
                    }
                })
            else:
                return jsonify({
                    "configured": True,
                    "connected": False,
                    "message": f"Connection failed: {response.status_code} - {response.text}",
                    "error": response.text
                })
                
        except Exception as e:
            return jsonify({
                "configured": True,
                "connected": False,
                "message": f"Connection error: {str(e)}",
                "error": str(e)
            })
            
    except Exception as e:
        return jsonify({
            "configured": False,
            "connected": False,
            "message": f"Error checking status: {str(e)}",
            "error": str(e)
        })

@app.route("/api/ado/test", methods=['POST'])
def test_ado_creation():
    """Test Azure DevOps work item creation"""
    try:
        # Check if Azure DevOps is configured
        if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
            return jsonify({
                "success": False,
                "message": "Azure DevOps not configured"
            })
        
        # Create a test work item
        test_title = "BA Agent Test Work Item"
        test_description = "This is a test work item created by BA Agent to verify Azure DevOps integration."
        
        result = create_ado_work_item("Task", test_title, test_description, tags=["BA-Agent-Test"])
        
        if result:
            return jsonify({
                "success": True,
                "message": "Test work item created successfully",
                "work_item": {
                    "id": result.get('id'),
                    "url": result.get('url'),
                    "title": test_title
                }
            })
        else:
            return jsonify({
                "success": False,
                "message": "Failed to create test work item"
            })
            
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"Error testing Azure DevOps: {str(e)}"
        })

if __name__ == '__main__':
    # Avoid Flask re-parsing .env (we already handled loading in config)
    app.run(debug=True, port=5000, load_dotenv=False)
